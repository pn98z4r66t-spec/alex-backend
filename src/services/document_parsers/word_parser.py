"""
Word Document Parser
Extracts text, tables, and metadata from Word documents (.docx)
"""
from docx import Document
from typing import Dict, List, Any
import logging

logger = logging.getLogger(__name__)


class WordParser:
    """Parser for Word documents (.docx)"""
    
    @staticmethod
    def parse(file_path: str) -> Dict[str, Any]:
        """
        Parse Word document and extract content
        
        Args:
            file_path: Path to .docx file
            
        Returns:
            Dictionary containing:
                - text: Full text content
                - paragraphs: List of paragraphs
                - tables: Extracted table data
                - metadata: Document properties
                - has_images: Whether document contains images
        """
        try:
            result = {
                'text': '',
                'paragraphs': [],
                'tables': [],
                'metadata': {},
                'has_images': False,
                'success': True,
                'error': None
            }
            
            doc = Document(file_path)
            
            # Extract core properties (metadata)
            core_props = doc.core_properties
            result['metadata'] = {
                'title': core_props.title or '',
                'author': core_props.author or '',
                'subject': core_props.subject or '',
                'keywords': core_props.keywords or '',
                'created': str(core_props.created) if core_props.created else '',
                'modified': str(core_props.modified) if core_props.modified else '',
                'last_modified_by': core_props.last_modified_by or '',
            }
            
            # Extract paragraphs
            all_text = []
            for para in doc.paragraphs:
                if para.text.strip():
                    result['paragraphs'].append({
                        'text': para.text,
                        'style': para.style.name if para.style else 'Normal'
                    })
                    all_text.append(para.text)
            
            # Extract tables
            for table_idx, table in enumerate(doc.tables, 1):
                table_data = {
                    'table_number': table_idx,
                    'rows': len(table.rows),
                    'columns': len(table.columns),
                    'data': []
                }
                
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data['data'].append(row_data)
                
                result['tables'].append(table_data)
                
                # Add table content to text
                table_text = '\n'.join([' | '.join(row) for row in table_data['data']])
                all_text.append(f"\n[Table {table_idx}]\n{table_text}\n")
            
            # Check for images
            for rel in doc.part.rels.values():
                if "image" in rel.target_ref:
                    result['has_images'] = True
                    break
            
            # Combine all text
            result['text'] = '\n\n'.join(all_text)
            
            # Add statistics
            result['statistics'] = {
                'total_characters': len(result['text']),
                'total_words': len(result['text'].split()),
                'paragraph_count': len(result['paragraphs']),
                'table_count': len(result['tables']),
            }
            
            return result
            
        except Exception as e:
            logger.error(f"Error parsing Word document: {e}")
            return {
                'success': False,
                'error': str(e),
                'text': '',
                'paragraphs': [],
                'tables': [],
                'metadata': {}
            }
    
    @staticmethod
    def get_summary(parsed_data: Dict[str, Any]) -> str:
        """
        Generate a human-readable summary of the Word document
        
        Args:
            parsed_data: Output from parse() method
            
        Returns:
            Summary string
        """
        if not parsed_data.get('success'):
            return f"Failed to parse Word document: {parsed_data.get('error')}"
        
        summary_parts = []
        
        # Basic info
        stats = parsed_data.get('statistics', {})
        summary_parts.append(
            f"Word Document with {stats.get('paragraph_count', 0)} paragraphs"
        )
        
        # Metadata
        metadata = parsed_data.get('metadata', {})
        if metadata.get('title'):
            summary_parts.append(f"Title: {metadata['title']}")
        if metadata.get('author'):
            summary_parts.append(f"Author: {metadata['author']}")
        
        # Content statistics
        summary_parts.append(
            f"Content: {stats.get('total_words', 0)} words, "
            f"{stats.get('total_characters', 0)} characters"
        )
        
        # Tables
        if stats.get('table_count', 0) > 0:
            summary_parts.append(f"Contains {stats['table_count']} table(s)")
        
        # Images
        if parsed_data.get('has_images'):
            summary_parts.append("Contains images")
        
        return '\n'.join(summary_parts)

